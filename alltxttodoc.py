# file per convertire tutti i file txt in una cartella in docx
from docx import Document
import glob
import os

doc = Document()

testo = []
for txt in glob.glob("*txt"):
    with open(txt) as file:
        file = file.read()
        testo.append((txt[:-4], file))

for t in testo:
    doc.add_heading(t[0], 1)
    doc.add_paragraph(t[1])

doc.save("word.doc")

os.system("start word.doc")
